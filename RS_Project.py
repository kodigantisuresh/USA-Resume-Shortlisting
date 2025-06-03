import os
import subprocess
import pandas as pd
import re
import time
import logging
import io
import json
import hashlib
from datetime import datetime, timedelta
from dateutil.parser import parse
from google.oauth2.credentials import Credentials
from google.auth.transport.requests import Request
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from base64 import urlsafe_b64decode
from tabulate import tabulate
from docx import Document
from PyPDF2 import PdfReader
import shutil
import tempfile
import atexit
import pytz
import google.generativeai as genai
from pdfminer.high_level import extract_text
from datetime import datetime

# Configure logging
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

# Configuration
script_dir = os.getcwd()
RESUME_FOLDER = os.path.join(script_dir, "Resumes")
LIBREOFFICE_PATH = r"C:\\Program Files\\LibreOffice\\program\\soffice.exe"
OUTPUT_CSV = os.path.join(RESUME_FOLDER, "resume_analysis.csv")
GEMINI_API_KEY = "AIzaSyD9pk2mrmj_POX9KoouBwbUXKAdrIWaCqA"
CURRENT_DATE = datetime.now().strftime("%Y-%m-%d")
SKILLS_LIST = []

# Initialize Gemini
genai.configure(api_key=GEMINI_API_KEY)
available_models = [m.name for m in genai.list_models()]
if 'models/gemini-1.5-pro-latest' in available_models:
    model = genai.GenerativeModel('models/gemini-1.5-pro-latest')
elif 'models/gemini-pro' in available_models:
    model = genai.GenerativeModel('models/gemini-pro')
else:
    raise ValueError("No supported Gemini model available")

def is_technical_skill(skill):
    """Determine if a skill is technical based on enhanced heuristics."""
    skill_lower = skill.lower().strip()
    if re.search(r'efrain|pacheco|job id|va-\d+|richmond', skill_lower):
        return False
    if any(suffix in skill_lower for suffix in ['.js', '.net', '.py', '.java', '.php', '.sql']):
        return True
    words = skill_lower.split()
    if len(words) > 1 and all(word[0].isupper() for word in words if word):
        return True
    if len(words) == 1 and len(skill_lower) >= 2:
        if skill_lower.endswith(('script', 'lang')):
            return True
    if len(words) > 1:
        if any(char in skill_lower for char in ['+', '#', '.']):
            return True
        if any(word.startswith(('cyber', 'dev', 'net', 'web', 'data', 'cloud', 'micro', 'server', 'auto')) for word in words):
            return True
    if len(words) == 1 and len(skill_lower) < 2:
        return False
    if re.search(r'[a-zA-Z]+\\d+|\\d+[a-zA-Z]+', skill_lower):
        return True
    if len(words) == 1 and len(skill_lower) >= 3:
        if re.match(r'^[A-Z][a-z]+$', skill_lower) or skill_lower.endswith(('sql', 'js', 'py', 'java')):
            return True
    return False

'''def extract_skills_from_subject(subject):
    """Extracts technical skills from the email subject with improved tokenization."""
    logging.info(f"Extracting skills from subject: {subject}")
    subject = re.sub(r'^\s*Fwd:\s*', '', subject, flags=re.IGNORECASE)
    subject = re.sub(r'\b(Hybrid|Local|Remote|Onsite|Certified|\d+\+?|Rate:.*?\d+\.?\d*|Job ID:.*?|VA-\d+|Richmond|W2)\b[\s/]*', '', subject, flags=re.IGNORECASE)
    subject = re.sub(r'\([^()]*\)', '', subject)
    subject = re.sub(r'\s{2,}', ' ', subject)
    subject = subject.strip(',; |')
    logging.info(f"Cleaned subject: {subject}")
    skills = []
    separators = r'\s*(?:,|\sand\s|\swith\s|;|/|\|)\s*'
    def process_phrase(phrase, prefix=""):
        phrase = phrase.strip()
        if not phrase:
            return
        nested_separators = r'\s*(?:/|\|)\s*'
        if re.search(nested_separators, phrase):
            parts = re.split(nested_separators, phrase)
            first_part_words = parts[0].split()
            local_prefix = prefix
            remaining_parts = parts
            if len(first_part_words) > 1 and not prefix:
                local_prefix = first_part_words[0]
                remaining_parts[0] = " ".join(first_part_words[1:])
            elif len(first_part_words) == 1 and not prefix:
                local_prefix = first_part_words[0]
                remaining_parts = parts[1:] if len(parts) > 1 else []
            for part in remaining_parts:
                part = part.strip()
                if not part:
                    continue
                combined = f"{local_prefix} {part}".strip()
                if is_technical_skill(combined):
                    skills.append(combined.lower())
                if is_technical_skill(part):
                    skills.append(part.lower())
                process_phrase(part, local_prefix)
            return
        words = phrase.split()
        if not words:
            return
        i = 0
        while i < len(words):
            found_skill = False
            for j in range(len(words), i, -1):
                potential_skill = " ".join(words[i:j]).strip()
                if prefix:
                    combined_skill = f"{prefix} {potential_skill}".strip()
                    if is_technical_skill(combined_skill):
                        skills.append(combined_skill.lower())
                        i = j
                        found_skill = True
                        break
                if is_technical_skill(potential_skill):
                    skills.append(potential_skill.lower())
                    i = j
                    found_skill = True
                    break
            if not found_skill:
                word = words[i]
                if prefix:
                    combined_skill = f"{prefix} {word}".strip()
                    if is_technical_skill(combined_skill):
                        skills.append(combined_skill.lower())
                        i += 1
                        continue
                if is_technical_skill(word):
                    skills.append(word.lower())
                i += 1
    phrases = re.split(separators, subject, flags=re.IGNORECASE)
    phrases = [phrase.strip() for phrase in phrases if phrase.strip()]
    logging.info(f"Phrases after splitting: {phrases}")
    for phrase in phrases:
        process_phrase(phrase)
    seen = set()
    skills = [s for s in skills if not (s in seen or seen.add(s))]
    logging.info(f"Extracted skills: {skills}")
    return skills'''

def auto_authenticate_google():
    """Automatically authenticates with Google APIs."""
    creds = None
    token_path = 'token.json'
    SCOPES = ['https://www.googleapis.com/auth/gmail.readonly']
    if os.path.exists(token_path):
        try:
            creds = Credentials.from_authorized_user_file(token_path, SCOPES)
            if creds.expired and creds.refresh_token:
                creds.refresh(Request())
        except Exception as e:
            logging.error(f"Error loading credentials: {e}")
            creds = None
    if not creds or not creds.valid:
        try:
            flow = InstalledAppFlow.from_client_secrets_file('client.json', SCOPES)
            creds = flow.run_local_server(port=8080)
            token_data = json.loads(creds.to_json())
            token_data['creation_time'] = datetime.now(pytz.UTC).isoformat()
            with open(token_path, 'w') as token:
                json.dump(token_data, token)
        except Exception as e:
            logging.error(f"Authentication failed: {e}")
            raise
    try:
        gmail_service = build('gmail', 'v1', credentials=creds)
        return gmail_service
    except Exception as e:
        logging.error(f"Failed to build services: {e}")
        raise

def get_emails_by_job_id(service, job_id):
    """Fetch all emails containing a given job ID within the last 30 days."""
    today = datetime.now()
    last_month = today - timedelta(days=30)
    last_month_str = last_month.strftime('%Y/%m/%d')
    all_messages = []
    page_token = None
    while True:
        try:
            results = service.users().messages().list(
                userId="me",
                q=f"after:{last_month_str} {job_id}",
                maxResults=500,
                pageToken=page_token
            ).execute()
            messages = results.get("messages", [])
            all_messages.extend(messages)
            page_token = results.get("nextPageToken")
            if not page_token:
                break
            time.sleep(1)
        except Exception as e:
            logging.error(f"Error fetching emails: {e}")
            break
    return all_messages

def decode_base64(data):
    """Decodes base64 email content safely."""
    missing_padding = len(data) % 4
    if missing_padding:
        data += '=' * (4 - missing_padding)
    return urlsafe_b64decode(data)

def extract_email_body(payload):
    """Extracts the email body content."""
    if not payload:
        return ""
    if "body" in payload and "data" in payload["body"]:
        return decode_base64(payload["body"]["data"]).decode("utf-8", errors="ignore")
    if "parts" in payload:
        for part in payload["parts"]:
            if part.get("mimeType", "") in ["text/plain", "text/html"] and "data" in part.get("body", {}):
                return decode_base64(part["body"]["data"]).decode("utf-8", errors="ignore")
            if "parts" in part:
                nested_body = extract_email_body(part)
                if nested_body:
                    return nested_body
    return ""

def extract_details_from_body(body):
    """Extracts comprehensive candidate details from email body."""
    name_patterns = [
        r"First\s*Name\s*\(per DL\):\s*([^\n]+)\s*Middle\s*Name\s*\(.*?\):\s*([^\n]*)\s*Last\s*Name\s*\(per DL\):\s*([^\n]+)(?=\s*(?:\n|Year\s+of\s+Birth|$))",
        r"First\s*Name\s*\(.*?\):\s*(.*?)\s*Middle\s*Name\s*\(.*?\):\s*(.*?)\s*Last\s*Name\s*\(.*?\):\s*(.*?)(?=\s*(?:\n|Year\s+of\s+Birth|$))",
        r"Name\s*\(per\s*DL\):\s*([^\n]+)\s*\n\s*Last\s*Name\s*\(per\s*DL BOD\):\s*([^\n]+)(?=\s*(?:\n|Year\s+of\s+Birth|$))",
        r"First\s*Name\s*:\s*([^\n]+)\s*Middle\s*Name\s*\(.*?\):\s*(?:[^\n]*)\s*Last\s*Name\s*:\s*([^\n]+)(?=\s*(?:\n|Year\s+of\s+Birth|$))",
        r"Name\s*:\s*(.*?)(?=\s*(?:\n|Year\s+of\s+Birth|$))",
        r"<(b|strong)>(.*?)</\1>(?=(?:[^<]*(?:<(?!\/?b|strong)[^>]*>))*?(?:\n|Year\s+of\s+Birth|$))",
    ]
    name = "N/A"
    for pattern in name_patterns:
        match = re.search(pattern, body, re.IGNORECASE | re.DOTALL)
        if match:
            groups = [g.strip() for g in match.groups() if g]
            if len(groups) >= 3 and 'per DL' in pattern:
                first = groups[0]
                middle = groups[1] if groups[1] and not groups[1].lower().startswith('blank') else ''
                last = groups[2]
                name = ' '.join(filter(None, [first, middle, last]))
                break
            cleaned_groups = []
            for g in groups:
                cleaned = re.sub(r'\s*>\s*', ' ', g)
                cleaned = re.sub(r'[\*><\(\)\s]+', ' ', cleaned).strip()
                cleaned_groups.append(cleaned)
            name = " ".join(filter(None, cleaned_groups))
            break
    location_match = re.search(
        r"(?i)Current\s+location\s*(?:\s*\(city/state\))?\s*[:\-]?\s*([^\n]+?)\s*(?=\n|$)",
        body
    )
    current_location = location_match.group(1).strip() if location_match else "N/A"
    yob_match = re.search(r"Year\s*of\s*Birth\s*(?:\(19xx\))?\s*[:\-]?\s*(\d{2,4})", body, re.IGNORECASE)
    yob = yob_match.group(1) if yob_match else "N/A"
    if yob and len(yob) == 2:
        yob = f"19{yob}" if int(yob) > 24 else f"20{yob}"
    cert_match = re.search(r"Certification\s*Count\s*[:\-]?\s*(NA|\d+)", body, re.IGNORECASE)
    certification_count = 0
    if cert_match:
        cert_str = cert_match.group(1).strip()
        if cert_str.isdigit():
            certification_count = int(cert_str)
    visa_patterns = [
        r"Visa\s*Status\s*with\s*Validity\s*:\s*([^\n\r]*)",
        r"Visa\s*type\s*and\s*sponsor\s*name\s*\(.*?\)\s*:\s*([^\n\r]*)",
        r"Visa\s*type\s*:\s*([^\n\r]*)",
        r"Status\s*:\s*([^\n\r]*)"
    ]
    visa_info = "N/A"
    for pattern in visa_patterns:
        match = re.search(pattern, body, re.IGNORECASE)
        if match:
            raw_visa_info = match.group(1).strip()
            visa_info = re.sub(r'[\*\s]+', ' ', raw_visa_info).strip()
            if not visa_info or visa_info.isspace():
                visa_info = "N/A"
            break
    skills_section = re.search(r'(?:Skills|Technical Skills|Key Skills|Proficiencies)\s*[:\n](.*?)(?=\n\n|\Z)', body, re.IGNORECASE | re.DOTALL)
    skills = []
    if skills_section:
        skills_text = skills_section.group(1).strip()
        skills = [s.strip() for s in re.split(r'[,\n;]+', skills_text) if s.strip() and is_technical_skill(s.strip())]
        skills = list(dict.fromkeys(skills))
    return {
        "Name": name,
        "Year of Birth": yob,
        "Current Location": current_location,
        "Certification Count": certification_count,
        "Visa Status": visa_info.strip(),
        "Skills": skills
    }

def fetch_attachments(service, message_id):
    """Fetch all attachments from an email."""
    try:
        msg = service.users().messages().get(userId="me", id=message_id, format="full").execute()
        payload = msg.get("payload", {})
        parts = payload.get("parts", [])
        attachments = []
        for part in parts:
            filename = part.get("filename")
            attachment_id = part.get("body", {}).get("attachmentId")
            if filename and attachment_id:
                attachments.append((filename, attachment_id))
        logging.info(f"Fetched {len(attachments)} attachments for message {message_id}")
        return attachments
    except Exception as e:
        logging.error(f"Error fetching attachments for message {message_id}: {e}")
        return []

def get_attachment_data(service, message_id, attachment_id):
    """Fetch and decode attachment data using the Gmail API."""
    try:
        attachment = service.users().messages().attachments().get(
            userId="me",
            messageId=message_id,
            id=attachment_id
        ).execute()
        data = attachment.get("data", "")
        return urlsafe_b64decode(data)
    except Exception as e:
        logging.error(f"Error fetching attachment data for attachment {attachment_id}: {e}")
        return None

def extract_text_from_attachment(attachment_data, filename, temp_dir=None):
    """Extract text from attachment (PDF, DOCX, DOC) with fallback."""
    try:
        if filename.lower().endswith(".docx"):
            doc = Document(io.BytesIO(attachment_data))
            return "\n".join(para.text for para in doc.paragraphs if para.text.strip())
        elif filename.lower().endswith(".pdf"):
            try:
                text = extract_text(io.BytesIO(attachment_data))
                if text.strip():
                    return text
            except Exception as e:
                logging.warning(f"pdfminer failed for {filename}: {e}")
            # Fallback to PyPDF2
            try:
                reader = PdfReader(io.BytesIO(attachment_data))
                text = ""
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                return text
            except Exception as e:
                logging.warning(f"PyPDF2 failed for {filename}: {e}")
                return None
        elif filename.lower().endswith(".doc"):
            if not temp_dir:
                temp_dir = tempfile.mkdtemp()
            temp_doc_path = os.path.join(temp_dir, "temp.doc")
            with open(temp_doc_path, "wb") as f:
                f.write(attachment_data)
            for attempt in range(2):  # Retry once
                try:
                    docx_path = convert_doc_to_docx(temp_doc_path, temp_dir)
                    if docx_path.endswith(".docx"):
                        doc = Document(docx_path)
                        text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
                        os.remove(docx_path)
                        os.remove(temp_doc_path)
                        return text
                    os.remove(temp_doc_path)
                    return None
                except Exception as e:
                    logging.warning(f"Attempt {attempt+1} failed for {filename}: {e}")
                    time.sleep(1)  # Brief pause before retry
            os.remove(temp_doc_path)
            return None
        return None
    except Exception as e:
        logging.error(f"Error extracting text from {filename}: {e}")
        return None
    finally:
        if temp_dir and os.path.exists(temp_dir):
            shutil.rmtree(temp_dir, ignore_errors=True)

def filter_excluded_files(attachments):
    """Filter out non-resume attachments based on extensions and keywords."""
    EXCLUDED_EXTENSIONS = [".jpeg", ".jpg", ".png", ".gif", ".bmp", ".xlsx"]
    EXCLUDED_KEYPHRASES = [
        "dl", "visa", "h1", "gc", "i-129", "approval",
    "sm", "skill matrix", "rtr", "innosoul",
    "reference", "patibandla", "check form", "sow", "ead",
    "70125071", "scanned", "driver license", "driving license",
    "passport", "i9", "w2", "paystub", "offer letter", "contract",
    "background check", "ssn", "social security", "id card",
    "certification form", "certification_form", "self-certification",
    "certification checklist", "certification doc", "cert doc",
    "certification sheet", "cert sheet", "certification template",
    "certification application", "cert app", "certification request",
    "cert req", "certification verification", "cert verify",
    "certification status", "cert status", "certification record",
    "cert record", "certification notice", "cert notice",
    "certification confirmation", "cert confirm", "certification letter",
    "cert letter", "certification card", "cert card",
    "certification badge", "cert badge", "certification id",
    "cert id", "certification proof", "cert proof",
    "F-Resume-Self-Certification-Form.pdf", "self-certification",
    "authorization form", "approval form", "clearance form",
    "verification form", "compliance form", "disclosure form",
    "attestation form", "acknowledgment form", "agreement form",
    "consent form", "declaration form", "enrollment form",
    "registration form", "application form", "submission form",
    "request form", "approval form", "clearance certificate",
    "security clearance", "background form", "screening form"
    ]
    valid_files = []
    for filename, attachment_id in attachments:
        lower_name = filename.lower()
        if any(lower_name.endswith(ext) for ext in EXCLUDED_EXTENSIONS):
            logging.debug(f"Skipping non-resume file (extension): {filename}")
            continue
        if any(term in lower_name for term in EXCLUDED_KEYPHRASES):
            logging.debug(f"Skipping non-resume file (keyword): {filename}")
            continue
        if lower_name.endswith((".pdf", ".docx", ".doc")):
            valid_files.append((filename, attachment_id))
    logging.info(f"Filtered {len(attachments)} attachments to {len(valid_files)} valid files")
    return valid_files

def is_resume_content(text):
    """Validate if text represents a resume with keyword scoring."""
    if not text:
        return False
    NON_RESUME_PATTERNS = [
        r"reference\s*check", r"visa\s*status", 
        r"approval\s*notice", r"skill\s*matrix",
        r"return\s*to\s*recruiter", r"form\s*[0-9]{3}",
        r"government\s*issued", r"validity\s*date",
        r"solicitation\s*number", r"candidate\s*reference"
    ]
    RESUME_SECTIONS = [
        r"work\s*experience", r"professional\s*(history|experience|summary)",
        r"skills?", r"education", r"projects?", r"certifications?",
        r"technical\s*(skills|proficiencies)", r"employment\s*history",
        r"key\s*skills", r"executive\s*summary", r"work\s*history",
        r"technical\s*summary", r"professional\s*overview"
    ]
    POSITIVE_KEYWORDS = [
        r"\d+\+?\s*years?\s*of\s*experience", r"developed", r"implemented",
        r"power\s*apps", r"power\s*automate", r"power\s*bi", r"azure",
        r"dataverse", r"microsoft\s*365", r"dynamics\s*365", r"sharepoint",
        r"certified", r"bachelor", r"master", r"engineer", r"developer",
        r"architect", r"solution", r"automation", r"integration"
    ]
    text = re.sub(r"\s+", " ", text.lower())
    # Check for non-resume patterns
    if any(re.search(pattern, text) for pattern in NON_RESUME_PATTERNS):
        logging.debug(f"Non-resume content detected in text: {text[:100]}...")
        return False
    # Calculate resume score
    section_count = sum(1 for pattern in RESUME_SECTIONS if re.search(pattern, text))
    keyword_count = sum(1 for pattern in POSITIVE_KEYWORDS if re.search(pattern, text))
    score = section_count * 2 + keyword_count
    is_resume = score >= 3  # Require a score of 3 (e.g., 1 section + 1 keyword or 3 keywords)
    logging.debug(f"Resume validation: sections={section_count}, keywords={keyword_count}, score={score}, is_resume={is_resume}")
    return is_resume

def validate_resume(service, message_id, attachment_id, filename):
    """Validate if attachment is a resume by checking its content."""
    # First check against excluded files
    if not is_potential_resume(filename):
        return False
    
    # Get attachment data
    data = get_attachment_data(service, message_id, attachment_id)
    if not data:
        return False
    
    # Extract text
    text = extract_text_from_attachment(data, filename)
    if not text:
        return False
    
    # Check for resume content
    return is_resume_content(text)

def is_potential_resume(filename):
    """Check if filename could potentially be a resume."""
    lower_name = filename.lower()
    
    # Excluded patterns
    EXCLUDED_KEYPHRASES = [
        "dl", "visa", "h1", "gc", "i-129", "approval",
    "sm", "skill matrix", "rtr", "innosoul",
    "reference", "patibandla", "check form", "sow", "ead",
    "70125071", "scanned", "driver license", "driving license",
    "passport", "i9", "w2", "paystub", "offer letter", "contract",
    "background check", "ssn", "social security", "id card",
    "certification form", "certification_form", "self-certification",
    "certification checklist", "certification doc", "cert doc",
    "certification sheet", "cert sheet", "certification template",
    "certification application", "cert app", "certification request",
    "cert req", "certification verification", "cert verify",
    "certification status", "cert status", "certification record",
    "cert record", "certification notice", "cert notice",
    "certification confirmation", "cert confirm", "certification letter",
    "cert letter", "certification card", "cert card",
    "certification badge", "cert badge", "certification id",
    "cert id", "certification proof", "cert proof",
    "F-Resume-Self-Certification-Form.pdf", "self-certification",
    "authorization form", "approval form", "clearance form",
    "verification form", "compliance form", "disclosure form",
    "attestation form", "acknowledgment form", "agreement form",
    "consent form", "declaration form", "enrollment form",
    "registration form", "application form", "submission form",
    "request form", "approval form", "clearance certificate",
    "security clearance", "background form", "screening form"
    ]
    
    # Check against excluded patterns
    if any(term in lower_name for term in EXCLUDED_KEYPHRASES):
        return False
    
    # Check for valid extensions
    if not lower_name.endswith((".pdf", ".docx", ".doc")):
        return False
    
    return True
def identify_resume(service, message_id, attachments):
    """
    Enhanced Resume Identification with 3-Phase Approach:
    1. Priority check for clearly named resumes (resume, cv, name-based)
    2. Content validation of remaining files
    3. Fallback to most likely candidate if no clear resume found
    
    Returns filename of the identified resume or "N/A" if none found
    """
    # Phase 1: Priority check for obvious resumes
    priority_candidates = []
    other_candidates = []
    
    for filename, attachment_id in attachments:
        lower_name = filename.lower()
        
        # Check for standard resume filename patterns
        is_standard_resume = (
            "resume" in lower_name or 
            "cv" in lower_name or 
            "curriculum vitae" in lower_name or
            "bio data" in lower_name
        )
        
        # Check for name-based patterns (e.g., "JohnDoe.pdf", "Smith_Profile.docx")
        name_pattern = (
            re.match(r"^[A-Z][a-z]+[A-Z][a-z]+\.(pdf|docx|doc)$", filename) or  # JohnDoe.pdf
            re.match(r"^[A-Z][a-z]+_[A-Z][a-z]+\.(pdf|docx|doc)$", filename) or  # John_Doe.pdf
            "profile" in lower_name or
            "portfolio" in lower_name
        )
        
        # Check for valid extensions
        valid_extension = lower_name.endswith((".pdf", ".docx", ".doc"))
        
        if (is_standard_resume or name_pattern) and valid_extension:
            priority_candidates.append((filename, attachment_id))
        elif valid_extension:
            other_candidates.append((filename, attachment_id))
    
    # Check priority candidates first
    for filename, attachment_id in priority_candidates:
        if validate_resume(service, message_id, attachment_id, filename):
            logging.info(f"Identified resume by filename: {filename}")
            return filename
    
    # Phase 2: Validate remaining files by content
    validated_resumes = []
    for filename, attachment_id in other_candidates:
        if validate_resume(service, message_id, attachment_id, filename):
            validated_resumes.append(filename)
    
    # Handle multiple potential resumes
    if len(validated_resumes) == 1:
        logging.info(f"Identified resume by content: {validated_resumes[0]}")
        return validated_resumes[0]
    elif len(validated_resumes) > 1:
        # If multiple candidates, return the most resume-like one
        best_candidate = None
        highest_score = 0
        
        for filename in validated_resumes:
            # Get content to score
            attachment_id = next(aid for (fn, aid) in other_candidates if fn == filename)
            data = get_attachment_data(service, message_id, attachment_id)
            text = extract_text_from_attachment(data, filename)
            
            if not text:
                continue
                
            # Score based on resume sections found
            score = 0
            resume_sections = [
                r"work\s*experience", r"professional\s*(history|experience|summary)",
                r"skills?", r"education", r"projects?", r"certifications?",
                r"technical\s*(skills|proficiencies)", r"employment\s*history",
                r"summary\s*of\s*qualifications", r"career\s*objective"
            ]
            
            for section in resume_sections:
                if re.search(section, text, re.IGNORECASE):
                    score += 1
            
            if score > highest_score:
                highest_score = score
                best_candidate = filename
        
        if best_candidate:
            logging.info(f"Selected best resume from multiple candidates: {best_candidate}")
            return best_candidate
        return validated_resumes[0]  # Fallback to first if scoring failed
    
    # Phase 3: Fallback - if no clear resume found, return the first document that's not excluded
    if other_candidates:
        logging.warning(f"No clear resume found, returning first candidate: {other_candidates[0][0]}")
        return other_candidates[0][0]
    
    logging.warning("No valid resume found in attachments")
    return "N/A"


def save_resumes_to_folder(service, details, message_id, attachments, resume_folder):
    """Save the identified resume to the specified folder."""
    resume_filename = details["Resume File"]
    if resume_filename == "N/A":
        logging.info(f"No resume file found for candidate: {details.get('Name', 'Unknown')}")
        return
    # Normalize filename to avoid special characters
    safe_filename = re.sub(r'[^\w\s.-]', '_', resume_filename)
    safe_filename = re.sub(r'\s+', '_', safe_filename)  # Replace spaces with underscores
    attachment_id = next((aid for (fn, aid) in attachments if fn == resume_filename), None)
    if not attachment_id:
        logging.warning(f"Attachment ID not found for resume file: {resume_filename}")
        return
    attachment_data = get_attachment_data(service, message_id, attachment_id)
    if not attachment_data:
        logging.error(f"Failed to fetch attachment data for resume file: {resume_filename}")
        return
    resume_file_path = os.path.join(resume_folder, safe_filename)
    try:
        with open(resume_file_path, "wb") as file:
            file.write(attachment_data)
        logging.info(f"Saved resume file: {resume_file_path}")
    except Exception as e:
        logging.error(f"Failed to save resume file {safe_filename}: {e}")

def create_resume_folder(folder_name="Resumes"):
    """Creates a folder to store resumes, clearing existing files."""
    if os.path.exists(folder_name):
        for filename in os.listdir(folder_name):
            file_path = os.path.join(folder_name, filename)
            try:
                if os.path.isfile(file_path) or os.path.islink(file_path):
                    os.unlink(file_path)
                elif os.path.isdir(file_path):
                    shutil.rmtree(file_path)
            except Exception as e:
                logging.error(f"Failed to delete {file_path}: {e}")
    else:
        os.makedirs(folder_name)
    return folder_name

def convert_doc_to_docx(doc_path, output_folder):
    """Convert DOC to DOCX using LibreOffice."""
    try:
        subprocess.run([
            LIBREOFFICE_PATH,
            "--headless",
            "--convert-to", "docx",
            "--outdir", output_folder,
            doc_path
        ], check=True, timeout=30)
        docx_file = os.path.join(output_folder, os.path.splitext(os.path.basename(doc_path))[0] + ".docx")
        if os.path.exists(docx_file):
            logging.info(f"Converted {doc_path} to {docx_file}")
            return docx_file
        logging.warning(f"DOCX conversion failed for {doc_path}: Output file not found")
        return doc_path
    except subprocess.CalledProcessError as e:
        logging.error(f"Error converting {doc_path} to DOCX: {e}")
        return doc_path
    except subprocess.TimeoutExpired:
        logging.error(f"Timeout converting {doc_path} to DOCX")
        return doc_path


def convert_docx_to_txt(input_path, output_path):
    """Convert .docx to .txt, extracting text from all possible document elements."""
    try:
        doc = Document(input_path)
        full_text = []

        def extract_paragraphs(paragraphs):
            for para in paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())

        extract_paragraphs(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        full_text.append(cell_text)
                    for nested_table in cell.tables:
                        for nested_row in nested_table.rows:
                            for nested_cell in nested_row.cells:
                                nested_text = nested_cell.text.strip()
                                if nested_text:
                                    full_text.append(nested_text)

        for section in doc.sections:
            for header in section.header.paragraphs:
                if header.text.strip():
                    full_text.append(header.text.strip())
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text.strip())
            for footer in section.footer.paragraphs:
                if footer.text.strip():
                    full_text.append(footer.text.strip())
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text.strip())

        for shape in doc.element.body.iter():
            if shape.tag.endswith('wps:txbx'):
                for p in shape.findall('.//w:p', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                    text = ''.join(t.text for t in p.findall('.//w:t', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}) if t.text)
                    if text.strip():
                        full_text.append(text.strip())

        seen = set()
        full_text = [t for t in full_text if not (t in seen or seen.add(t))]

        if not full_text:
            logging.warning(f"No text extracted from {input_path}")
            full_text = ["No text content found in document"]

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(full_text))

        file_size = os.path.getsize(output_path)
        if file_size < 100:
            logging.warning(f"Generated {output_path} is very small ({file_size} bytes), may not contain full text")

        logging.info(f"Successfully converted {input_path} to {output_path} ({file_size} bytes)")
    except Exception as e:
        logging.error(f"Error converting {input_path} to .txt: {e}")
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(f"Error extracting text: {str(e)}")


def convert_pdf_to_txt(input_path, output_path):
    """Convert PDF to text using pdfminer with PyPDF2 fallback."""
    try:
        text = extract_text(input_path)
        if text.strip():
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
            logging.info(f"Successfully converted {input_path} to {output_path}")
            return
    except Exception as e:
        logging.warning(f"pdfminer failed for {input_path}: {e}")
    try:
        reader = PdfReader(input_path)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)
        logging.info(f"Successfully converted {input_path} to {output_path} using PyPDF2")
    except Exception as e:
        logging.error(f"Error converting {input_path} to .txt: {e}")
        raise

'''def extract_experience_with_date_patterns(resume_text):
    """Extract work experience, dates, and metrics from resume text."""
    resume_text = re.sub(r'[\t\r\n]+', ' ', resume_text)
    resume_text = re.sub(r'[–—]+', '-', resume_text)
    resume_text = re.sub(r'\s+', ' ', resume_text.strip())
    date_patterns = [
        r'(\d{1,2}/\d{2,4})\s*[-–—]\s*(\d{1,2}/\d{2,4}|Present|Current|Till Date|Now)',
        r'(\b\w+\s+\d{4})\s*[-–—]\s*(\b\w+\s+\d{4}|Present|Current|Till Date|Now)',
        r'(\d{4})\s*[-–—]\s*(\d{4}|Present|Current|Till Date|Now)',
        r'(\d{1,2}-\d{4})\s*[-–—]\s*(\d{1,2}-\d{4}|Present|Current|Till Date|Now)',
        r'(\b\w+\s+\d{2,4})\s*to\s*(\b\w+\s+\d{2,4}|Present|Current|Till Date|Now)',
        r'(\d{1,2}/\d{2,4})\s*to\s*(\d{1,2}/\d{2,4}|Present|Current|Till Date|Now)',
        r'(\d{4}-\d{1,2})\s*[-–—]\s*(\d{4}-\d{1,2}|Present|Current|Till Date|Now)',
        r'(\b\w+\s+\d{4})\s*-\s*(\d{4})',
        r'(Since\s+\d{4})',
        r'(\d{4})\s*[-–—]\s*(\d{4}-\d{1,2})'
    ]
    periods = []
    for pattern in date_patterns:
        matches = re.finditer(pattern, resume_text, re.IGNORECASE)
        for match in matches:
            start, end = match.groups() if len(match.groups()) == 2 else (match.group(1), None)
            try:
                if start.lower().startswith('since'):
                    start_date = parse(start[5:].strip(), fuzzy=True, default=datetime(2000, 1, 1))
                    start_date = start_date.replace(day=1)
                else:
                    start_date = parse(start, fuzzy=True, default=datetime(2000, 1, 1))
                    start_date = start_date.replace(day=1)
                if len(start.split('/')[-1]) == 2:
                    year = int(start.split('/')[-1])
                    start_date = start_date.replace(year=2000 + year if year <= 24 else 1900 + year)
                elif len(start.split('-')[-1]) == 2:
                    year = int(start.split('-')[-1])
                    start_date = start_date.replace(year=2000 + year if year <= 24 else 1900 + year)
                if not end or end.lower() in ['present', 'current', 'till date', 'now']:
                    end_date = datetime.strptime(CURRENT_DATE, '%Y-%m-%d')
                else:
                    end_date = parse(end, fuzzy=True, default=datetime(2000, 12, 31))
                    end_date = end_date.replace(day=28)
                    if len(end.split('/')[-1]) == 2:
                        year = int(end.split('/')[-1])
                        end_date = end_date.replace(year=2000 + year if year <= 24 else 1900 + year)
                    elif len(end.split('-')[-1]) == 2:
                        year = int(end.split('-')[-1])
                        end_date = end_date.replace(year=2000 + year if year <= 24 else 1900 + year)
                if end_date >= start_date:
                    periods.append((start_date, end_date))
                else:
                    logging.warning(f"Invalid period (end before start): {start} to {end}")
            except ValueError as e:
                logging.warning(f"Failed to parse dates {start} to {end}: {e}")
                continue
    if not periods:
        logging.info("No valid experience periods found")
        return {
            "total_experience_years": 0.0,
            "confidence": 0.5,
            "periods": [],
            "timeline": {},
            "validation_metrics": {"overlapping_periods": 0, "total_unique_months": 0}
        }
    periods.sort(key=lambda x: x[0])
    total_days = 0
    unique_months = set()
    overlapping_periods = 0
    for i, (start, end) in enumerate(periods):
        duration = (end - start).days + 1
        total_days += duration
        current = start
        while current <= end:
            month_key = (current.year, current.month)
            unique_months.add(month_key)
            current = current.replace(day=1)
            if current.month == 12:
                current = current.replace(year=current.year + 1, month=1)
            else:
                current = current.replace(month=current.month + 1)
        for j, (other_start, other_end) in enumerate(periods):
            if i < j and start <= other_end and end >= other_start:
                overlap_days = (min(end, other_end) - max(start, other_start)).days + 1
                if overlap_days > 30:
                    overlapping_periods += 1
    total_experience_years = total_days / 365.25
    total_unique_months = len(unique_months)
    confidence = 0.9
    if overlapping_periods > 0:
        confidence -= 0.15 * (overlapping_periods / max(1, len(periods)))
    if total_unique_months < 12:
        confidence -= 0.1
    if total_experience_years < 1:
        confidence -= 0.1
    confidence = max(0.6, min(confidence, 0.95))
    earliest_start = min(start for start, _ in periods).strftime('%Y-%m-%d') if periods else None
    latest_end = max(end for _, end in periods).strftime('%Y-%m-%d') if periods else None
    result = {
        "total_experience_years": round(total_experience_years, 2),
        "confidence": round(confidence, 2),
        "periods": [(start.strftime('%Y-%m-%d'), end.strftime('%Y-%m-%d')) for start, end in periods],
        "timeline": {
            "earliest_start": earliest_start,
            "latest_end": latest_end
        },
        "validation_metrics": {
            "overlapping_periods": overlapping_periods,
            "total_unique_months": total_unique_months
        }
    }
    logging.info(f"Experience extraction result: {result}")
    return result'''

def extract_work_experience_from_folder(folder_path, api_key):
    """Extract and calculate total work experience from all .txt resume files."""
    results = []
    for filename in os.listdir(folder_path):
        if filename.endswith('.txt'):
            file_path = os.path.join(folder_path, filename)
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    resume_text = file.read()
                prompt = f"""
                From the following resume text, extract all work experiences including:
                - Job title
                - Company name
                - Start date (month and year, e.g., 'January 2020', 'Jan 2020', '01/2020')
                - End date (month and year, 'Present', or 'Current')
                
                Return the result in a structured JSON format like:
                [
                    {{
                        "job_title": "Software Engineer",
                        "company": "Tech Corp",
                        "start_date": "January 2020",
                        "end_date": "March 2023"
                    }},
                    ...
                ]
                
                If no work experience is found or dates are missing, return an empty list: []
                
                Resume text:
                {resume_text}
                """
                response = model.generate_content(prompt)
                response_text = response.text.strip()
                response_text = re.sub(r'^```json\s*|\s*```$', '', response_text, flags=re.MULTILINE)
                try:
                    experiences = json.loads(response_text)
                    if not isinstance(experiences, list):
                        logging.error(f"Invalid JSON format for {filename}: Expected a list, got {type(experiences)}")
                        results.append({'Filename': filename, 'Experience': '0.00 years'})
                        continue
                except json.JSONDecodeError as json_err:
                    logging.error(f"Error parsing JSON for {filename}: {json_err}")
                    results.append({'Filename': filename, 'Experience': '0.00 years'})
                    continue
                date_ranges = []
                for exp in experiences:
                    start_date = exp.get('start_date')
                    end_date = exp.get('end_date')
                    if not start_date or start_date.lower() in ['none', '']:
                        logging.warning(f"Invalid start date in {filename}: {start_date}")
                        continue
                    try:
                        start = parse(start_date, fuzzy=True, default=datetime(2000, 1, 1))
                        start = datetime(start.year, start.month, 1)
                    except (ValueError, TypeError) as e:
                        logging.warning(f"Failed to parse start date in {filename}: {start_date}, Error: {e}")
                        continue
                    if not end_date or end_date.lower() in ['none', '']:
                        logging.warning(f"Invalid end date in {filename}: {end_date}")
                        continue
                    if end_date.lower() in ['present', 'current']:
                        end = datetime.now()
                    else:
                        try:
                            end = parse(end_date, fuzzy=True, default=datetime(2000, 1, 1))
                            end = datetime(end.year, end.month, 1)
                        except (ValueError, TypeError) as e:
                            logging.warning(f"Failed to parse end date in {filename}: {end_date}, Error: {e}")
                            continue
                    if end >= start:
                        date_ranges.append((start, end))
                    else:
                        logging.warning(f"Invalid date range in {filename}: {start_date} to {end_date}")
                date_ranges.sort(key=lambda x: x[0])
                merged_ranges = []
                if date_ranges:
                    current_start, current_end = date_ranges[0]
                    for start, end in date_ranges[1:]:
                        if start <= current_end:
                            current_end = max(current_end, end)
                        else:
                            merged_ranges.append((current_start, current_end))
                            current_start, current_end = start, end
                    merged_ranges.append((current_start, current_end))
                total_months = 0
                for start, end in merged_ranges:
                    delta = (end.year - start.year) * 12 + end.month - start.month
                    if delta > 0:
                        total_months += delta
                    elif delta == 0:
                        total_months += 1
                experience = total_months / 12
                experience = f"{experience:.2f} years"
                results.append({'Filename': filename, 'Experience': experience})
            except Exception as e:
                logging.error(f"Error processing {filename}: {e}")
                results.append({'Filename': filename, 'Experience': '0.00 years'})
    return pd.DataFrame(results)

def extract_candidate_details_with_gemini(resume_text):
    """Extracts candidate details from resume text using Gemini API."""
    prompt = f"""
    Extract the following details from the resume text below:
    1. Full Name (combine first, middle if available, and last names)
    2. Year of Birth (look for birth year, typically in format YYYY)
    3. Current Location (city and state/country)
    4. Visa Status (look for terms like H1B, GC, USC, OPT, etc.)
    
    Return the result in JSON format:
    {{
        "name": "Full Name",
        "year_of_birth": "YYYY",
        "current_location": "City, State/Country",
        "visa_status": "Visa Type"
    }}
    
    If any information is not found, use "N/A" as the value.
    
    Resume text:
    {resume_text}
    """
    try:
        response = model.generate_content(prompt)
        response_text = response.text.strip()
        response_text = re.sub(r'^```json\s*|\s*```$', '', response_text, flags=re.MULTILINE)
        details = json.loads(response_text)
        return {
            "Name": details.get("name", "N/A"),
            "Year of Birth": details.get("year_of_birth", "N/A"),
            "Current Location": details.get("current_location", "N/A"),
            "Visa Status": details.get("visa_status", "N/A")
        }
    except Exception as e:
        logging.error(f"Error extracting details with Gemini: {e}")
        return {
            "Name": "N/A",
            "Year of Birth": "N/A",
            "Current Location": "N/A",
            "Visa Status": "N/A"
        }

def extract_resume_details(folder_path, api_key, subject_skills):
    """Extract certification count, government work, matched skills and candidate details from .txt resume files."""
    subject_skills_lower = [skill.lower() for skill in subject_skills]
    results = []
    skill_frequency = {}
   
    for filename in os.listdir(folder_path):
        if filename.endswith('.txt'):
            file_path = os.path.join(folder_path, filename)
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    resume_text = file.read()
 
                # Extract candidate details using Gemini
                candidate_details = extract_candidate_details_with_gemini(resume_text)
               
                # Initialize matched skills dictionary
                matched_skills = {skill.lower(): 0 for skill in subject_skills}
               
                # Count exact matches (case-insensitive) for each skill in the resume text
                for skill in subject_skills:
                    # Create regex pattern that matches the whole word (to avoid partial matches)
                    pattern = re.compile(r'\b' + re.escape(skill.lower()) + r'\b', re.IGNORECASE)
                    matches = pattern.findall(resume_text.lower())
                    count = len(matches)
                   
                    if count > 0:
                        matched_skills[skill.lower()] = count
                        skill_frequency[skill.lower()] = skill_frequency.get(skill.lower(), 0) + count
 
                # Filter out skills with zero counts
                matched_skills = {k: v for k, v in matched_skills.items() if v > 0}
 
                # Extract certification and government work
                prompt = f"""
                From the following resume text, extract:
                1. Certification Count: Count all mentioned certifications
                2. Government Work: Check for USA government work experience
               
                Return in JSON format:
                {{
                    "certification_count": int,
                    "government_work": {{
                        "worked_with_govt": bool,
                        "govt_entities": list
                    }}
                }}
               
                Resume text:
                {resume_text}
                """
               
                response = model.generate_content(prompt)
                response_text = response.text.strip()
                response_text = re.sub(r'^```json\s*|\s*```$', '', response_text, flags=re.MULTILINE)
               
                try:
                    data = json.loads(response_text)
                except json.JSONDecodeError:
                    data = {
                        'certification_count': 0,
                        'government_work': {'worked_with_govt': False, 'govt_entities': []}
                    }
 
                results.append({
                    'Filename': filename,
                    'Name': candidate_details['Name'],
                    'Year of Birth': candidate_details['Year of Birth'],
                    'Current Location': candidate_details['Current Location'],
                    'Visa Status': candidate_details['Visa Status'],
                    'Certification Count': data.get('certification_count', 0),
                    'Government Work': 'Yes: ' + ', '.join(data['government_work']['govt_entities'])
                                    if data['government_work']['worked_with_govt'] else 'No',
                    'Matched Skills': matched_skills  # Store as dict to preserve counts
                })
 
            except Exception as e:
                logging.error(f"Error processing {filename}: {e}")
                results.append({
                    'Filename': filename,
                    'Name': "N/A",
                    'Year of Birth': "N/A",
                    'Current Location': "N/A",
                    'Visa Status': "N/A",
                    'Certification Count': 0,
                    'Government Work': 'No',
                    'Matched Skills': {}
                })
 
    if skill_frequency:
        freq_df = pd.DataFrame.from_dict(skill_frequency, orient='index', columns=['Count'])
        freq_df = freq_df.sort_values('Count', ascending=False)
    else:
        logging.info("No skills matched across resumes")
   
    return pd.DataFrame(results)

def format_skills_comma_separated(matched_skills):
    """Format matched skills dictionary as a comma-separated string with counts."""
    if not matched_skills:
        return "None"
    return ", ".join([f"{skill} ({count})" for skill, count in matched_skills.items() if count > 0]) or "None"

def process_folder(folder_path):
    """Convert PDF, DOC, and DOCX files to TXT with robust error handling."""
    processed_files = set()
    filename_mapping = {}  # Track original to .txt filename mappings
    for filename in os.listdir(folder_path):
        if filename in processed_files:
            continue
        file_path = os.path.join(folder_path, filename)
        base_name = os.path.splitext(filename)[0]
        # Normalize base_name for consistency
        safe_base_name = re.sub(r'[^\w\s.-]', '_', base_name)
        safe_base_name = re.sub(r'\s+', '_', safe_base_name)
        txt_path = os.path.join(folder_path, f"{safe_base_name}.txt")
        processed_files.add(filename)
        filename_mapping[filename] = os.path.basename(txt_path)
        try:
            if filename.lower().endswith('.pdf'):
                convert_pdf_to_txt(file_path, txt_path)
                if os.path.exists(txt_path):
                    os.remove(file_path)
                else:
                    logging.warning(f"No text file created for {filename}")
            elif filename.lower().endswith('.doc'):
                docx_path = convert_doc_to_docx(file_path, folder_path)
                if docx_path.endswith('.docx'):
                    text = extract_text_from_attachment(open(docx_path, 'rb').read(), docx_path)
                    if text:
                        with open(txt_path, 'w', encoding='utf-8') as f:
                            f.write(text)
                        logging.info(f"Converted {docx_path} to {txt_path}")
                    if docx_path != file_path:
                        os.remove(docx_path)
                os.remove(file_path)
                if not os.path.exists(txt_path):
                    logging.warning(f"No text file created for {filename}")
            elif filename.lower().endswith('.docx'):
                convert_docx_to_txt(file_path, txt_path)
                if os.path.exists(txt_path):
                    os.remove(file_path)
                    logging.info(f"Converted {file_path} to {txt_path} using convert_docx_to_txt")
                else:
                    logging.warning(f"No text file created for {filename}")
        except Exception as e:
            logging.error(f"Error processing {filename}: {e}")
            continue
    txt_files = [f for f in os.listdir(folder_path) if f.endswith('.txt')]
    logging.info(f"Processed folder, found {len(txt_files)} .txt files: {txt_files}")
    logging.info(f"Filename mapping: {filename_mapping}")
    return filename_mapping

def apply_resume_scoring(df):
    """Applies scenario-based ranking based on Matched Skills, Government, and Experience."""
    logging.info(f"Applying resume scoring. DataFrame columns: {list(df.columns)}")
    required_columns = ["Matched Skills", "Experience", "Government Work"]
    missing_columns = [col for col in required_columns if col not in df.columns]
    if missing_columns:
        logging.error(f"Missing required columns: {missing_columns}")
        raise KeyError(f"Missing required columns: {missing_columns}")
    def calculate_matched_skills_count(matched_skills):
        if pd.isna(matched_skills) or matched_skills == "None":
            return 0
        try:
            skills = matched_skills.split(", ")
            total_count = 0
            for skill in skills:
                match = re.search(r'\((\d+)\)', skill)
                if match:
                    total_count += int(match.group(1))
            return total_count
        except Exception as e:
            logging.warning(f"Error parsing matched skills '{matched_skills}': {e}")
            return 0
    def parse_experience(experience_str):
        if pd.isna(experience_str) or not isinstance(experience_str, str) or experience_str == "0 years":
            return 0.0
        match = re.match(r"(\d+\.?\d*)\s*years?", experience_str)
        if match:
            return float(match.group(1))
        return 0.0
    def calculate_government_count(government_experience):
        if (
            pd.notna(government_experience)
            and isinstance(government_experience, str)
            and government_experience.strip().lower() not in ["no", "no govt", "not worked with the government"]
            and government_experience.strip() != ""
        ):
            if government_experience.lower() == "yes":
                return 1
            return len([entity for entity in government_experience.split(", ") if entity.strip() and entity.lower() != "yes"])
        return 0
    def detect_scenario(df):
        skills_present = any(df["Matched Skills Count"] > 0)
        gov_present = any(df["Government Count"] > 0)
        if not skills_present and gov_present:
            return "Scenario 1: Prioritize Skills (70%), Government (20%), Experience (10%)"
        elif not gov_present and skills_present:
            return "Scenario 2: Prioritize Skills (80%), Experience (20%)"
        elif skills_present and gov_present:
            return "Scenario 3: Prioritize Government (50%), Experience (50%)"
        return "Scenario 3: Prioritize Government (50%), Experience (50%)"
    df["Matched Skills Count"] = df["Matched Skills"].apply(calculate_matched_skills_count)
    df["Experience Years"] = df["Experience"].apply(parse_experience)
    df["Government Count"] = df["Government Work"].apply(calculate_government_count)
    scenario = detect_scenario(df)
    logging.info(f"Detected Scenario: {scenario}")
    if "Scenario 1" in scenario:

        max_skills = df["Matched Skills Count"].max() if df["Matched Skills Count"].max() > 0 else 1
        max_gov = df["Government Count"].max() if df["Government Count"].max() > 0 else 1
        max_exp = df["Experience Years"].max() if df["Experience Years"].max() > 0 else 1
        df["Composite Score"] = (
            0.7 * (df["Matched Skills Count"] / max_skills) +
            0.2 * (df["Government Count"] / max_gov) +
            0.1 * (df["Experience Years"] / max_exp)
        )
    elif "Scenario 2" in scenario:
        max_skills = df["Matched Skills Count"].max() if df["Matched Skills Count"].max() > 0 else 1
        max_exp = df["Experience Years"].max() if df["Experience Years"].max() > 0 else 1
        df["Composite Score"] = (
            0.8 * (df["Matched Skills Count"] / max_skills) +
            0.2 * (df["Experience Years"] / max_exp)
        )
    else:
        max_gov = df["Government Count"].max() if df["Government Count"].max() > 0 else 1
        max_exp = df["Experience Years"].max() if df["Experience Years"].max() > 0 else 1
        df["Composite Score"] = (
            0.5 * (df["Government Count"] / max_gov) +
            0.5 * (df["Experience Years"] / max_exp)
        )
    df["Rank"] = df["Composite Score"].rank(ascending=False, method="min").astype(int)
    df = df.drop(columns=["Composite Score", "Experience Years", "Government Count"], errors="ignore")
    df = df.sort_values(by="Rank", ascending=True).reset_index(drop=True)
    return df, scenario


'''def extract_subject_skills(subject):
    """Extracts skills from subject by separating job role and skills."""
    if not subject:
        return []

    # Clean the subject
    subject = re.sub(r'^\s*(?:Fwd:|Re:)\s*', '', subject, flags=re.IGNORECASE)
    subject = re.sub(r'\([^)]*\)', '', subject)  # Remove parenthetical content
    subject = re.sub(r'\s{2,}', ' ', subject).strip()

    # Find where the job role likely ends (before first skill separator)
    role_end_match = re.search(r'[,/]', subject)
    role_end_pos = role_end_match.start() if role_end_match else len(subject)
    
    # Extract skills from remaining text after job role
    skills_text = subject[role_end_pos + 1:].strip() if role_end_match else ""
    
    # Split skills by common separators
    skills = []
    if skills_text:
        skill_chunks = re.split(r'\s*[,/]\s*|\s+and\s+|\s*-\s*', skills_text)
        for chunk in skill_chunks:
            chunk = chunk.strip()
            if chunk and len(chunk) > 1:  # Minimum 2 characters
                skills.append(chunk)

    # Remove duplicates while preserving order
    seen = set()
    unique_skills = []
    for skill in skills:
        lower_skill = skill.lower()
        if lower_skill not in seen:
            seen.add(lower_skill)
            unique_skills.append(skill)

    return unique_skills'''
def extract_job_role_and_skills(subject):
    """Cleanly separates job role name and subject skills without prefixes."""
    if not subject:
        return "", []
    
    # Clean the subject line
    subject = re.sub(r'^\s*(?:Fwd:|Re:)\s*', '', subject, flags=re.IGNORECASE)
    subject = re.sub(r'\([^)]*\)', '', subject)  # Remove parenthetical content
    subject = re.sub(r'\s{2,}', ' ', subject).strip()
    
    # Remove all location/employment type prefixes and suffixes aggressively
    prefixes = r'\b(?:Remote|Hybrid|Local|Onsite|Contract|Full[\s-]*Time|Part[\s-]*Time|W2|C2C)\b'
    subject = re.sub(
        fr'{prefixes}[\s\-:/,]*|[,\s\-:/]*{prefixes}(?=\s|$|[,/])',
        '', 
        subject, 
        flags=re.IGNORECASE
    )
    
    # Find where skills listing begins (after job role)
    # Updated regex to avoid matching hyphens within job roles
    skill_start = re.search(
        r'\b(?:with|using|needs?|requires?|looking\s+for|skills?|exp?|experience|proficient\s+in)\b\s*[:,/]?|\s*[,/]\s*(?=[A-Z0-9])', 
        subject, 
        re.IGNORECASE
    )
    
    if skill_start:
        # Split into job role and skills
        job_role = subject[:skill_start.start()].strip()
        skills_text = subject[skill_start.end():].strip()
        # If the match was a comma or slash, include it in skills_text
        if skill_start.group(0).strip() in [',', '/']:
            skills_text = skill_start.group(0).strip() + ' ' + skills_text
    else:
        # No clear transition, assume everything is job role
        job_role = subject
        skills_text = ""
    
    # Further clean job role, preserving specific characters like . in .NET
    job_role = re.sub(r'^[^a-zA-Z0-9.]+|[^a-zA-Z0-9.]+$', '', job_role)  # Allow . in addition to alphanumeric
    job_role = re.sub(r'\s+', ' ', job_role).strip()
    
    # Additional cleaning to remove any stray prefixes within the job role
    job_role = re.sub(fr'^{prefixes}\s*|\s*{prefixes}\s*$', '', job_role, flags=re.IGNORECASE)
    
    # Extract skills from skills text
    skills = []
    if skills_text:
        # Split by common separators while preserving multi-word skills
        skill_chunks = re.findall(r'(?:[^,/]|/(?=\s))+', skills_text)
        for chunk in skill_chunks:
            chunk = re.sub(r'^\W+|\W+$', '', chunk.strip())  # Clean each skill
            if chunk and len(chunk) > 1:  # Minimum 2 characters
                skills.append(chunk)
    
    return job_role, skills

def extract_email_data(service, message_id):
    """Extract candidate details, skills, job role and resume from email."""
    try:
        msg = service.users().messages().get(userId="me", id=message_id, format="full").execute()
        payload = msg.get("payload", {})
        headers = payload.get("headers", [])
        subject = next((header["value"] for header in headers if header["name"].lower() == "subject"), "")
        
        # Extract job role and skills from subject
        job_role, subject_skills = extract_job_role_and_skills(subject)
        
        body = extract_email_body(payload)
        details = extract_details_from_body(body)
        attachments = fetch_attachments(service, message_id)
        resume_filename = identify_resume(service, message_id, attachments)
        details.update({
            "Resume File": resume_filename,
            "Job Role": job_role  # Add job role to details
        })
        skills = details.pop("Skills", [])
        return details, skills, subject_skills
    except Exception as e:
        logging.error(f"Error processing email {message_id}: {e}")
        return {}, [], []

# Updated main() function
def main(job_id):
    """Main function with enhanced resume handling, diagnostics, and deduplication."""
    try:
        service = auto_authenticate_google()
        logging.info(f"Starting processing for job ID: {job_id}")
        resume_folder = create_resume_folder(RESUME_FOLDER)
        messages = get_emails_by_job_id(service, job_id)
        if not messages:
            print(f"No emails found related to job ID {job_id}.")
            logging.info(f"No emails found for job ID {job_id}. Terminating.")
            return
        print(f"Found {len(messages)} emails related to job ID {job_id}:")
        logging.info(f"Processing {len(messages)} emails")
        extracted_skills_list = []
        email_data = []
        subject_skills_list = []
        failure_reasons = []
        for msg in messages:
            message_id = msg["id"]
            details, skills, subject_skills = extract_email_data(service, message_id)
            job_role = details.get("Job Role", "")
            if skills or subject_skills:
                subject_skills_list.append({
                    "Job Role": job_role,
                    "Subject": subject_skills,
                    "Skills": skills
                })
            else:
                logging.warning(f"No skills extracted from body for message {message_id}")
            if subject_skills:
                logging.info(f"Subject skills extracted for message {message_id}: {subject_skills}")
            else:
                logging.warning(f"No subject skills extracted for message {message_id}")
            if details:
                email_data.append(details)
                attachments = fetch_attachments(service, message_id)
                if attachments:
                    save_resumes_to_folder(service, details, message_id, attachments, resume_folder)
                else:
                    logging.info(f"No attachments found for message {message_id} (Candidate: {details.get('Name', 'Unknown')})")
                    failure_reasons.append(f"Message {message_id}: No attachments found for candidate {details.get('Name', 'Unknown')}")
            else:
                logging.warning(f"No candidate details extracted for message {message_id}")
                failure_reasons.append(f"Message {message_id}: Failed to extract candidate details")
        logging.info(f"Total candidates extracted: {len(email_data)}")
        logging.info(f"Total subject skills extracted: {len(subject_skills_list)}")
        # Filter subject_skills_list to include only entries with at least one technical skill
        filtered_subject_skills_list = [
            item for item in subject_skills_list
            if any(is_technical_skill(skill) for skill in item["Subject"])
        ]
        if filtered_subject_skills_list:
            print("\n=== Extracted Job Roles and Skills ===")
            for item in filtered_subject_skills_list:
                print(f"\nJob Role: {item.get('Job Role', 'N/A')}")
                print(f"Subject Skills: {', '.join(item['Subject']) if item['Subject'] else 'None'}")
                print("\nExtracted Skills from Body:")
                for skill in item['Skills']:
                    print(f"- {skill}")
                print("-" * 50)
        else:
            print("\nNo job roles with technical skills found.")
            logging.info("No job roles with technical skills found in subject_skills_list")
        if email_data:
            # Deduplicate email_data based on Name
            deduplicated_email_data = []
            name_to_entries = {}
            for entry in email_data:
                name = entry.get("Name", "N/A").strip().lower()
                if name == "n/a" or not name:
                    continue  # Skip invalid names
                if name not in name_to_entries:
                    name_to_entries[name] = []
                name_to_entries[name].append(entry)
            for name, entries in name_to_entries.items():
                if len(entries) == 1:
                    deduplicated_email_data.append(entries[0])
                else:
                    # Choose the entry with the most complete information
                    best_entry = max(entries, key=lambda x: (
                        x.get("Resume File", "N/A") != "N/A",  # Prioritize entries with a resume
                        len(x.get("Skills", [])),  # Then by number of skills
                        -sum(1 for v in x.values() if v == "N/A")  # Then by fewest N/A values
                    ))
                    deduplicated_email_data.append(best_entry)
                    logging.info(f"Deduplicated {len(entries)} entries for {name}, kept entry with resume: {best_entry.get('Resume File', 'N/A')}")
            logging.info(f"After deduplication, {len(deduplicated_email_data)} candidates remain")
            df = pd.DataFrame(deduplicated_email_data)
            logging.info(f"Created DataFrame with {len(df)} candidates, columns: {list(df.columns)}")
            global SKILLS_LIST
            SKILLS_LIST = filtered_subject_skills_list[0]["Subject"] if filtered_subject_skills_list else []
            logging.info(f"Skills list for resume processing: {SKILLS_LIST}")
            try:
                filename_mapping = process_folder(RESUME_FOLDER)
                logging.info("Completed file conversion to .txt")
                experience_df = extract_work_experience_from_folder(RESUME_FOLDER, GEMINI_API_KEY)
                details_df = extract_resume_details(RESUME_FOLDER, GEMINI_API_KEY, SKILLS_LIST)
                resume_df = pd.merge(experience_df, details_df, on="Filename", how="outer")
                resume_df['Earliest Start'] = 'N/A'
                resume_df['Latest End'] = 'N/A'
                resume_df['Overlapping Periods'] = 0
                resume_df['Confidence'] = '80%'
                logging.info(f"Processed {len(resume_df)} resumes, resume_df columns: {list(resume_df.columns)}")
                if resume_df.empty:
                    print("\nNo .txt files found for analysis.")
                    logging.warning("No .txt files found for analysis")
                    return
                # Normalize filenames for merge
                def normalize_filename(filename):
                    if pd.isna(filename) or filename == "N/A":
                        return "N/A"
                    base = os.path.splitext(filename)[0]
                    base = re.sub(r'[^\w\s.-]', '_', base)
                    base = re.sub(r'\s+', '_', base)
                    return f"{base}.txt"
                df["Resume File"] = df["Resume File"].apply(normalize_filename)
                resume_df["Filename"] = resume_df["Filename"].apply(lambda x: x if pd.notna(x) else "N/A")
                # Log filenames for debugging
                logging.info(f"Email Resume Files: {df['Resume File'].tolist()}")
                logging.info(f"Processed .txt Files: {resume_df['Filename'].tolist()}")
                logging.info(f"Filename Mapping: {filename_mapping}")
                # Merge with exact match
                df = pd.merge(
                    df,
                    resume_df,
                    left_on="Resume File",
                    right_on="Filename",
                    how="left"
                )
                # Fallback: Partial match on candidate name
                unmatched = df[df['Filename'].isna() & (df['Resume File'] != "N/A")]
                if not unmatched.empty:
                    logging.warning(f"Unmatched resume files: {unmatched['Resume File'].tolist()}")
                    for idx, row in unmatched.iterrows():
                        candidate_name = row['Name'].lower().replace(' ', '_')
                        for resume_filename in resume_df['Filename']:
                            if candidate_name in resume_filename.lower():
                                logging.info(f"Fallback match: {row['Resume File']} matched to {resume_filename} via name {candidate_name}")
                                df.loc[idx, resume_df.columns] = resume_df[resume_df['Filename'] == resume_filename].iloc[0]
                                break
                # Log unmatched records
                unmatched = df[df['Filename'].isna() & (df['Resume File'] != "N/A")]
                if not unmatched.empty:
                    logging.warning(f"Still unmatched after fallback: {unmatched['Resume File'].tolist()}")
                # Handle name columns
                name_cols = ['Name', 'Year of Birth', 'Current Location', 'Visa Status']
                for col in name_cols:
                    if f"{col}_x" in df.columns and f"{col}_y" in df.columns:
                        df[col] = df[f"{col}_x"].fillna(df[f"{col}_y"])
                        df = df.drop([f"{col}_x", f"{col}_y"], axis=1)
                    elif f"{col}_x" in df.columns:
                        df[col] = df[f"{col}_x"]
                        df = df.drop([f"{col}_x"], axis=1)
                    elif f"{col}_y" in df.columns:
                        df[col] = df[f"{col}_y"]
                        df = df.drop([f"{col}_y"], axis=1)
                    else:
                        df[col] = "N/A"
                # Certification count handling
                cert_cols = [c for c in df.columns if 'Certification' in c]
                if len(cert_cols) > 1:
                    df['Certification Count'] = df[cert_cols[0]].fillna(df[cert_cols[1]]).fillna(0).astype(int)
                    df = df.drop(cert_cols, axis=1)
                elif cert_cols:
                    df['Certification Count'] = df[cert_cols[0]].fillna(0).astype(int)
                    df = df.drop(cert_cols, axis=1)
                else:
                    df['Certification Count'] = 0
                # Fill missing values to avoid nan
                df['Experience'] = df['Experience'].fillna('0.00 years')
                df['Government Work'] = df['Government Work'].fillna('No')
                df['Matched Skills'] = df['Matched Skills'].fillna('None')
                df = df.drop(["Resume File", "Filename"], axis=1, errors="ignore")
                # Filter out rows where all values are "N/A" (job description emails)
                df = df[~((df['Name'] == 'N/A') &
                          (df['Year of Birth'] == 'N/A') &
                          (df['Current Location'] == 'N/A') &
                          (df['Visa Status'] == 'N/A'))]
                # Add Job Role and Subject Skills to DataFrame
                if filtered_subject_skills_list:
                    job_role = filtered_subject_skills_list[0].get('Job Role', 'N/A')
                    subject_skills = ', '.join(filtered_subject_skills_list[0]['Subject']) if filtered_subject_skills_list[0]['Subject'] else 'None'
                    df['Job Role'] = job_role
                    df['Subject Skills'] = subject_skills
                else:
                    df['Job Role'] = 'N/A'
                    df['Subject Skills'] = 'None'
                logging.info(f"Final DataFrame columns after merge: {list(df.columns)}")
                df, scenario = apply_resume_scoring(df)
                print(f"\nRanking Criteria: {scenario}")
                logging.info(f"Applied resume scoring with scenario: {scenario}")
                df.to_csv(OUTPUT_CSV, index=False)
                logging.info(f"Results saved to {OUTPUT_CSV}")
                columns_order = [
                    "Rank", "Name", "Year of Birth", "Current Location", "Visa Status",
                    "Experience", "Certification Count", "Government Work", "Matched Skills"
                ]
                columns_order = [col for col in columns_order if col in df.columns]
                print("\nCandidate Details (Sorted by Rank):")
                print(tabulate(df[columns_order], headers="keys", tablefmt="grid", showindex=False))
            except Exception as e:
                print(f"Failed to process resumes: {e}")
                logging.error(f"Failed to process resumes: {e}")
                return
        else:
            print("\nNo candidate data or resumes found for comparison.")
            logging.warning("No candidate data or resumes found for comparison")
            if failure_reasons:
                print("\nReasons for failure:")
                for reason in failure_reasons:
                    print(f"- {reason}")
            else:
                print("\nNo specific reasons logged for failure. Check the logs for more details.")
    except Exception as e:
        logging.error(f"Error in main execution: {e}")
        print(f"An error occurred: {e}")

if __name__ == "__main__":
    job_id = input("Enter the job ID to search for: ").strip()
    main(job_id)
 